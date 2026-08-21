import base64
from io import BytesIO
import unittest

from docx import Document
from docx.oxml import parse_xml
from docx.shared import Inches

from app.parser import parse_docx_bytes


PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)


class ParserTests(unittest.TestCase):
    def test_parse_table_format_and_duplicate_pair(self):
        data = self._docx_bytes(
            [
                {
                    "serial": "05",
                    "category": "G-Science",
                    "question": "কোন গ্রহকে লাল গ্রহ বলা হয়?",
                    "options": ["পৃথিবী", "মঙ্গল", "শুক্র", "বুধ"],
                    "explanation": "মঙ্গল গ্রহকে লাল গ্রহ বলা হয়।",
                    "answer": "Answer: B",
                },
                {
                    "serial": "18",
                    "category": "G-Science",
                    "question": "কোন গ্রহকে লাল গ্রহ বলা হয়?",
                    "options": ["পৃথিবী", "মঙ্গল", "শুক্র", "বুধ"],
                    "explanation": "মঙ্গল গ্রহের পৃষ্ঠ লালচে।",
                    "answer": "খ",
                },
            ]
        )

        parsed = parse_docx_bytes(data)

        self.assertEqual(parsed["total"], 2)
        self.assertEqual(parsed["duplicateCount"], 1)
        self.assertEqual(parsed["duplicatePairs"][0]["source"]["serial"], "05")
        self.assertEqual(parsed["duplicatePairs"][0]["repeat"]["serial"], "18")
        self.assertEqual(parsed["questions"][0]["answerLabel"], "খ")
        self.assertEqual(parsed["questions"][0]["answerText"], "মঙ্গল")

    def test_parse_equation_and_image_parts(self):
        document = Document()
        table = document.add_table(rows=8, cols=2)
        table.cell(0, 0).text = "01"
        table.cell(0, 1).text = "Maths"

        question_paragraph = table.cell(1, 0).paragraphs[0]
        question_paragraph._p.append(
            parse_xml(
                """
                <m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
                  <m:r><m:t>x</m:t></m:r>
                  <m:sSup>
                    <m:e><m:r><m:t>y</m:t></m:r></m:e>
                    <m:sup><m:r><m:t>2</m:t></m:r></m:sup>
                  </m:sSup>
                  <m:f>
                    <m:num><m:r><m:t>3</m:t></m:r></m:num>
                    <m:den><m:r><m:t>11</m:t></m:r></m:den>
                  </m:f>
                </m:oMath>
                """
            )
        )
        question_paragraph.add_run(" সমীকরণটি দেখুন")

        table.cell(2, 0).text = "ছবি"
        table.cell(2, 0).paragraphs[0].add_run().add_picture(BytesIO(PNG_1X1), width=Inches(0.12))
        table.cell(3, 0).text = "দুই"
        table.cell(4, 0).text = "তিন"
        table.cell(5, 0).text = "চার"
        table.cell(6, 0).text = "ব্যাখ্যা"
        table.cell(7, 0).text = "A"

        output = BytesIO()
        document.save(output)
        parsed = parse_docx_bytes(output.getvalue())
        question = parsed["questions"][0]

        self.assertIn("xy^2", question["question"])
        self.assertIn("<math", question["questionHtml"])
        self.assertIn("<msup>", question["questionHtml"])
        self.assertIn("<mfrac>", question["questionHtml"])
        self.assertIn("data:image/png;base64", question["options"][0]["html"])
        self.assertIn("image", [part["type"] for part in question["options"][0]["parts"]])

    def _docx_bytes(self, rows):
        document = Document()
        for item in rows:
            table = document.add_table(rows=8, cols=2)
            table.cell(0, 0).text = item["serial"]
            table.cell(0, 1).text = item["category"]
            table.cell(1, 0).text = item["question"]
            for offset, option in enumerate(item["options"], start=2):
                table.cell(offset, 0).text = option
            table.cell(6, 0).text = item["explanation"]
            table.cell(7, 0).text = item["answer"]

        output = BytesIO()
        document.save(output)
        return output.getvalue()


if __name__ == "__main__":
    unittest.main()
